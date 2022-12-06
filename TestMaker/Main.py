from tkinter import *
from tkinter import ttk
from tkinter import filedialog
from PIL import Image, ImageTk
import re
from Variants import *
import datetime
from calendar import monthrange
#pip install python-docx
#https://python-docx.readthedocs.io/en/latest/
import docx

class App:
    def __init__(self):
        self.variant=Variant()
        self.window = Tk()
        self.window.title("Making class/home work")
        size="860x"+str(max(self.variant.getCatNumbers()*10,190))
        self.window.geometry(size)
        self.frame = Frame(self.window)
        self.frame.grid()

        img1 = Image.open("ico/Rus.jpg")
        self.IcoRus = ImageTk.PhotoImage(img1)
        img2 = Image.open("ico/Eng.jpg")
        self.IcoEng = ImageTk.PhotoImage(img2)
        self.lang="En"

        self.lblTypeOfWork = Label(self.frame, text="Choose type of work")
        self.lblTypeOfWork.grid(column=0, row=0, padx=10)
        self.varTypeWork = IntVar()
        self.radioClass = Radiobutton(self.frame, text="Classwork", variable=self.varTypeWork, value=2, command=self.__saveFileDir)
        self.radioClass.grid(column=1, row=0)
        self.radioHome = Radiobutton(self.frame, text="Homework", variable=self.varTypeWork, value=1, command=self.__saveFileDir)
        self.radioHome.grid(column=2, row=0)
        self.radioClass.select()

        now = datetime.datetime.now()
        self.lblDataSelect = LabelFrame(self.frame, text='Data selection', labelanchor=N)
        self.lblDay = Label(self.lblDataSelect , text="DD")
        self.lblDay.grid(column=0, row=0)
        self.daySpin = ttk.Spinbox(self.lblDataSelect , from_=1, to_=31, width=5, command=self.__correctData)
        self.daySpin.grid(column=1, row=0)
        self.daySpin.set(now.day)
        self.lblMonth = Label(self.lblDataSelect , text="MM")
        self.lblMonth.grid(column=2, row=0)
        self.monthSpin = ttk.Spinbox(self.lblDataSelect , from_=1, to_=12, width=5, command=self.__correctData)
        self.monthSpin.grid(column=3, row=0)
        self.monthSpin.set(now.month)
        self.lblYear = Label(self.lblDataSelect , text="YY")
        self.lblYear.grid(column=4, row=0)
        self.yearSpin = ttk.Spinbox(self.lblDataSelect , from_=now.year, to_=now.year+10, width=5, command=self.__correctData)
        self.yearSpin.set(now.year)
        self.yearSpin.grid(column=5, row=0)
        self.lblDataSelect.grid(column=0, row=1, columnspan=5, padx=10)

        self.filenameDir="C:"
        self.lblSave = Label(self.frame, text="Where to save")
        self.lblSave.grid(column=0, row=2, padx=10)
        self.btnDirSave = Button(self.frame, text="...", command=self.__clickedSaving)
        self.btnDirSave.grid(column=1, row=2)
        self.lblSaveFile = Label(self.frame, text="Will be saved as...")
        self.lblSaveFile.grid(column=0, row=3, columnspan=7, padx=10)
        self.__saveFileDir()

        self.lblVariant = Label(self.frame, text="Choose key:")
        self.lblVariant.grid(column=1, row=5, padx=10)
        self.entryKey = Entry(self.frame)
        self.entryKey.grid(column=2, row=5)
        self.entryKey.insert(0,now.strftime("%A")+str(now.day)+str(now.month)+str(now.year%100))
        self.variantKey = 0

        self.lblSaveTypeSelect = LabelFrame(self.frame, text='Which type', labelanchor=N)
        self.varSaveWork = IntVar()
        self.checkWorkSave = Checkbutton(self.lblSaveTypeSelect, text='Work', variable=self.varSaveWork)
        self.varSaveWork.set(1)
        self.checkWorkSave.grid(column=0, row=0)
        self.varSaveAnswer = IntVar()
        self.checkAnswerSave = Checkbutton(self.lblSaveTypeSelect, text='Answer', variable=self.varSaveAnswer)
        self.varSaveAnswer.set(1)
        self.checkAnswerSave.grid(column=1, row=0)
        self.lblSaveTypeSelect.grid(column=0, row=5, padx=10)

        self.btnLang = Button(self.frame, image=self.IcoEng, command=self.__ChangeLang)
        self.btnLang.grid(column=6, row=5)
        self.btnSave = Button(self.frame, text="Create&Save", command=self.__SavingVariant)
        self.btnSave.grid(column=2, row=6, columnspan=2)

        self.lblTasks=[]
        self.SpinTasks=[]
        self.lblNumsTasks = LabelFrame(self.frame, text='Tasks numbers', labelanchor=N)
        for i in range(self.variant.getCatNumbers()):
            self.lblTasks.append(Label(self.lblNumsTasks, text=str(i+1)+" cat:", anchor=W))
            self.lblTasks[i].grid(column=0, row=i, padx=10)
            self.SpinTasks.append(ttk.Spinbox(self.lblNumsTasks , from_=0, to_=3, width=5))
            self.SpinTasks[i].set(1)
            self.SpinTasks[i].grid(column=1, row=i, padx=10)
        self.lblNumsTasks.grid(column=7, row=0, rowspan=self.variant.getCatNumbers(), padx=10)

        self.__ChangeLang()
        self.window.mainloop()
    def __SetVariantKey(self):
        keyNum=0
        for ch in self.entryKey.get():
            if(ch<'0'):
                ch='0'
            keyNum=(keyNum*10+ord(ch)-ord('0'))
            keyNum=keyNum % 1000000
        self.variantKey=keyNum
        self.variant.setVariantNum(keyNum)
        self.variant.setVariantLang(self.lang)

    def __ChangeLang(self):
        if(self.lang=="En"):
            self.lang="Ru"
            self.btnLang.configure(image=self.IcoRus)
            self.lblTypeOfWork["text"] = "Выберите тип работы"
            self.lblNumsTasks["text"]='Задач по темам'
            self.radioClass['text'] = "Классная работа"
            self.radioHome['text'] = "Домашнаяя работа"
            self.lblDataSelect['text'] = 'Выбор даты работы'
            self.lblSave['text']="Куда сохранить"
            self.lblVariant['text'] ="Ключ генерации:"
            self.lblSaveTypeSelect['text'] = 'Сохранить'
            self.checkWorkSave['text'] = 'Работу'
            self.checkAnswerSave['text'] = 'Ответы'
            self.btnSave['text']="Создать+Сохранить"
            for i in range(self.variant.getCatNumbers()):
                self.lblTasks[i]['text']=str(i + 1) + ". " +self.variant.getAllCategoriesNameRus()[i]+":"
        else:
            self.lang="En"
            self.btnLang.configure(image=self.IcoEng)
            self.lblTypeOfWork["text"] = "Choose type of work"
            self.lblNumsTasks["text"]='Tasks numbers'
            self.radioClass['text']="Classwork"
            self.radioHome['text']="Homework"
            self.lblDataSelect['text']='Data selection'
            self.lblSave['text']="Where to save"
            self.lblVariant['text'] ="Choose key:"
            self.lblSaveTypeSelect['text'] = 'Which type'
            self.checkWorkSave['text'] = 'Work'
            self.checkAnswerSave['text'] = 'Answer'
            self.btnSave['text']="Create&Save"
            for i in range(self.variant.getCatNumbers()):
                self.lblTasks[i]['text']=str(i + 1) + ". " +self.variant.getAllCategoriesName()[i]+":"
        self.__saveFileDir()
    def __SavingVariant(self):
        self.__SetVariantKey()
        typeOfWork=""
        if(self.varTypeWork.get()==1):
            typeOfWork = "Домашняя работа "
        else:
            typeOfWork = "Классная работа "

        docA = docx.Document()
        docW = docx.Document()
        if(self.varSaveWork.get()):
            section = docW.sections[0]
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.text = typeOfWork+"\t\t"+self.daySpin.get()+"/"+self.monthSpin.get()+"/"+str(int(self.yearSpin.get())%100)
            paragraph.style = docW.styles["Header"]
            foot = section.footer.paragraphs[0]
            foot.text = "\t\t" + str(self.variantKey)
            foot.style = docA.styles["Footer"]
        if (self.varSaveAnswer.get()):
            section = docA.sections[0]
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.text = typeOfWork + "\t\t" + self.daySpin.get() + "/" + self.monthSpin.get() + "/" + str(
                int(self.yearSpin.get()) % 100)
            paragraph.style = docA.styles["Header"]
            foot = section.footer.paragraphs[0]
            foot.text = "\t\t" + str(self.variantKey)
            foot.style = docA.styles["Footer"]

        for i in range(self.variant.getCatNumbers()):
            for j in range(int(self.SpinTasks[i].get())):
                self.variant.setVariantNum(self.variantKey+j)
                if (self.varSaveWork.get()):
                    docW.add_paragraph(self.variant.getQuestion(i+1), style='List Number')
                if (self.varSaveAnswer.get()):
                    docA.add_paragraph(self.variant.getAnswer(i+1), style='List Number')

        if (self.varSaveWork.get()):
            docW.save(self.filenameDir+"/"+self.filenameSave)
        if (self.varSaveAnswer.get()):
            asnswFileName = self.filenameSave
            asnswFileName = asnswFileName.rstrip('.doc')
            asnswFileName = asnswFileName + "_answer.doc"
            docA.save(self.filenameDir+"/"+asnswFileName)

    def __saveFileDir(self):
        self.filenameSave=""
        if(self.varTypeWork.get()==1):
            self.filenameSave = "Homework "
        else:
            self.filenameSave = "Classwork "
        self.filenameSave = self.filenameSave+"("+self.daySpin.get()+"_"+self.monthSpin.get()+"_"+self.yearSpin.get() + ").doc"
        place=self.filenameDir +"/"+ self.filenameSave
        countMax=40
        if(len(place)>countMax):
            place = self.filenameDir[:(countMax-len(self.filenameSave))]
            place = place + ".../" + self.filenameSave
        if(self.lang=="En"):
            self.lblSaveFile["text"]="Will be saved as: "+place
        else:
            self.lblSaveFile["text"] = "Будет сохранено как: " + place

    def __correctData(self):
        days = monthrange(int(self.yearSpin.get()), int(self.monthSpin.get()))[1]
        if(int(self.daySpin.get())>days):
            self.daySpin.set(days)
        self.__saveFileDir()

    def __clickedSaving(self):
        file = filedialog.askdirectory()
        if (file == ""):
            return False
        self.filenameDir = file
        self.__saveFileDir()

app=App()